import sys
import random
from docx import Document

def shuffle_paragraphs(input_file, output_file):
    # Load the Word document
    doc = Document(input_file)

    # Extract paragraphs
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip() != '']

    # Shuffle paragraphs
    random.shuffle(paragraphs)

    # Create a new Word document
    new_doc = Document()

    # Add the shuffled paragraphs to the new document
    for para in paragraphs:
        new_doc.add_paragraph(para)

    # Save the new document
    new_doc.save(output_file)

def main():
    if len(sys.argv) != 3:
        print("Usage: python script.py input_file.docx output_file.docx")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    shuffle_paragraphs(input_file, output_file)
    print(f"Shuffled Word document saved as '{output_file}'")

if __name__ == "__main__":
    main()

